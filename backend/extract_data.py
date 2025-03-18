import json
import docx
from docx.table import Table
from typing import Dict, List, Any, Union
import os
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_table_data(table: Table) -> List[Dict[str, str]]:
    """Extract data from a table into a list of dictionaries."""
    data = []
    
    # Get headers from first row
    if len(table.rows) < 2:
        return data
        
    headers = []
    for cell in table.rows[0].cells:
        headers.append(cell.text.strip())
    
    # Extract data from remaining rows
    for i in range(1, len(table.rows)):
        row_data = {}
        for j, cell in enumerate(table.rows[i].cells):
            if j < len(headers):
                row_data[headers[j]] = cell.text.strip()
        data.append(row_data)
    
    return data

def extract_content(doc) -> Dict[str, Any]:
    """Extract structured content from a Word document."""
    result = {
        "title": "",
        "sections": []
    }
    
    current_section = None
    current_subsection = None
    tables = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Identify paragraph style to determine if it's a title, heading, etc.
        style = paragraph.style.name
        
        # Assume first non-empty paragraph with certain styles is the title
        if not result["title"] and style.startswith(('Title', 'Heading 1')):
            result["title"] = text
            continue

        # Identify headers by their style
        if style.startswith('Heading 1'):
            # Start a new section
            current_section = {
                "title": text,
                "content": [],
                "subsections": [],
                "tables": []
            }
            current_subsection = None
            result["sections"].append(current_section)
        
        elif style.startswith('Heading 2') and current_section:
            # Start a new subsection
            current_subsection = {
                "title": text,
                "content": [],
                "tables": []
            }
            current_section["subsections"].append(current_subsection)
            
        else:
            # Regular paragraph - add to current section or subsection
            if current_subsection:
                current_subsection["content"].append(text)
            elif current_section:
                current_section["content"].append(text)
    
    # Extract tables and associate them with sections
    for i, table in enumerate(doc.tables):
        table_data = extract_table_data(table)
        # For simplicity, we'll add tables to the most recent section
        if result["sections"]:
            result["sections"][-1]["tables"].append(table_data)
    
    return result

def convert_word_to_json(file_path: str, output_path: str = None) -> Dict[str, Any]:
    """Convert a Word document to a structured JSON."""
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {"error": f"File not found: {file_path}"}

        # Check if it's a docx file
        if not file_path.lower().endswith('.docx'):
            logger.error(f"File is not a Word document: {file_path}")
            return {"error": f"File is not a Word document: {file_path}"}
            
        # Open the document
        doc = docx.Document(file_path)
        
        # Extract content
        data = extract_content(doc)
        
        # Save to file if output path is provided
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            logger.info(f"JSON data saved to {output_path}")
            
        return data
            
    except Exception as e:
        logger.exception(f"Error converting Word to JSON: {str(e)}")
        return {"error": str(e)}

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        logger.error("Usage: python extract_data.py <input_file> [output_file]")
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else "output.json"
    
    result = convert_word_to_json(input_file, output_file)
    if "error" in result:
        logger.error(result["error"])
        sys.exit(1)
    else:
        logger.info("Conversion successful") 